# Paragraph Translator
from io import BytesIO
# import gradio as gr
# Def_04 Docx file to translated_Docx file
#from transformers import MarianMTModel, MarianTokenizer
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import nltk
from nltk.tokenize import sent_tokenize
from nltk.tokenize import LineTokenizer
nltk.download('punkt')
import math
import torch
from docx import Document
from time import sleep
from tqdm import tqdm
import docx
import time
from tqdm import notebook

def DocxSplitting(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    # length_fullText=len(fullText)
    return fullText
     
def split(list_a, chunk_size):

  for i in range(0, len(list_a), chunk_size):
    yield list_a[i:i + chunk_size]


# mname = 'Helsinki-NLP/opus-mt-en-hi'
# tokenizer = MarianTokenizer.from_pretrained(mname)
# model = MarianMTModel.from_pretrained(mname)
# model.to(device)
#@st.cache
def btTranslator2(Texts,num,filename):
    if torch.cuda.is_available():  
      dev = "cuda"
    else:  
      dev = "cpu" 
    device = torch.device(dev)
    # print(device)      
    files=Document()
    
    a="Helsinki-NLP/opus-mt-en-ru"
    b="Helsinki-NLP/opus-mt-ru-fr"
    c="Helsinki-NLP/opus-mt-fr-en"
    # d="Helsinki-NLP/opus-mt-es-en"
    langs=[a,b,c]
    # text=para

    for i,lang in zip(notebook.tqdm(langs,desc="Book Processing... ",position=0, leave=True),langs):  
          time.sleep(0.01)
          # mname = '/content/drive/MyDrive/Transformers Models/opus-mt-en-hi-Trans Model'
          tokenizer = AutoTokenizer.from_pretrained(lang)
          model = AutoModelForSeq2SeqLM.from_pretrained(lang)
          model.to(device)
          lt = LineTokenizer()
          batch_size = 64
          
          for i,para in zip(notebook.tqdm(Texts,desc=f"Paras Translating...",position=0, leave=True),Texts):
              # para=para.text
              # paragraphs = lt.tokenize(para)  
              translated_paragraphs = []
              # half_file=Document()
              # for i,paragraph in zip(notebook.tqdm(paragraphs,desc="Par Translating...",position=0, leave=True),paragraphs):
              time.sleep(0.01)
              sentences = sent_tokenize(para)
              batches = math.ceil(len(sentences) / batch_size)     
              translated = []
              for i in range(batches):
                  sent_batch = sentences[i*batch_size:(i+1)*batch_size]
                  model_inputs = tokenizer(sent_batch, return_tensors="pt", padding=True, truncation=True, max_length=500).to(device)
                  with torch.no_grad():
                      translated_batch = model.generate(**model_inputs)
                      translated += translated_batch
                  translated = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]
                  translated_paragraphs += [" ".join(translated)]
                  if lang[-2:]=='en':
                    files.add_paragraph(translated)
                  # else 
          # translated_text = "\n".join(translated_paragraphs)
          #bigtext=translated_text
          # half_file.add_paragraph(bigtext)
          # half_file.save(f"/content/drive/MyDrive/Translated Books/Half_Translated/{given_name}_{lang[-5:]}_T.docx")        
    # files.add_paragraph(translated_text) 
    #files2save=files.save("Translated.docx")
    #files.save("Translated.docx")
    #binary_output = BytesIO()
    #f=files.save(binary_output)
    #f2=f.getvalue()
    return files.save(f"/content/drive/MyDrive/Translated Books/{filename}_{num}_Translated.docx")


def DocxTranslate(file,filename):
    chunk_size = 400 #400 paras = 10000 WORDS 
    # file=r"/content/drive/MyDrive/Raw Books/01 Wife Cant Escape.docx"
    file=file
    texts=list(split(DocxSplitting(file), chunk_size))
    # filename="01 Wife Cant Escape"
    filename=filename
    for i,num,words10000 in zip(notebook.tqdm(texts,desc="Book Divided... ",position=0, leave=True),range(len(texts)),texts):
        print(f"==============={num}/{len(texts)}==============")
        btTranslator2(words10000,num,filename)
